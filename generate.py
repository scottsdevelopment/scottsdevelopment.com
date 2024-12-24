from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.shared import Pt, Inches
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime


def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    Creates a hyperlink within a paragraph.
    'paragraph' is a paragraph object we are adding the hyperlink to.
    'url' is the URL (or mailto/tel link) to link to.
    'text' is the text displayed for the hyperlink.
    'color' defaults to blue ("0000FF"), and 'underline' defaults to True.

    Returns the hyperlink object.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')

    # Set link color
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        r_pr.append(c)

    # Underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        r_pr.append(u)

    w_t = OxmlElement('w:t')
    w_t.text = text

    new_run.append(r_pr)
    new_run.append(w_t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def generate_new_resume(resume_data):
    """
    Creates a new Word document with the resume content
    and saves it as PersonName.Date.docx.
    Returns the name of the saved file.
    """
    doc = Document()

    # 1) Adjust Document Margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # 2) HEADER & FOOTER
    header = section.header
    if not header.paragraphs:
        header_para = header.add_paragraph()
    else:
        header_para = header.paragraphs[0]
    header_para.style = doc.styles['Normal']
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header_para.clear()

    # Name | Phone (tel) | Email (mailto)
    name_run = header_para.add_run(resume_data['heading']['name'])
    name_run.bold = True
    name_run.font.size = Pt(11)  # Slightly smaller in header
    header_para.add_run(" | ")

    add_hyperlink(
        header_para,
        f"tel:{resume_data['heading']['phone']}",
        resume_data['heading']['phone']
    )
    header_para.add_run(" | ")

    add_hyperlink(
        header_para,
        f"mailto:{resume_data['heading']['email']}",
        resume_data['heading']['email']
    )

    # Footer
    footer = section.footer
    if not footer.paragraphs:
        footer_para = footer.add_paragraph()
    else:
        footer_para = footer.paragraphs[0]
    footer_para.style = doc.styles['Normal']
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_para.clear()

    add_hyperlink(footer_para, resume_data['heading']
                  ['social_media'], resume_data['heading']['social_media'])

    # 3) Define/Update Basic Styles
    # "Best practice" settings for a resume
    # Name (Heading 1): 20pt, Bold
    heading1_style = ensure_paragraph_style(
        doc,
        style_name='Resume Heading 1',
        font_name='Calibri',
        font_size=20,
        bold=True,
        spacing_before=6,
        spacing_after=6
    )
    # Heading 2: 14pt, Bold
    heading2_style = ensure_paragraph_style(
        doc,
        style_name='Resume Heading 2',
        font_name='Calibri',
        font_size=14,
        bold=True,
        spacing_before=10,
        spacing_after=4
    )
    # Heading 3: 12pt, Bold
    heading3_style = ensure_paragraph_style(
        doc,
        style_name='Resume Heading 3',
        font_name='Calibri',
        font_size=12,
        bold=True,
        spacing_before=6,
        spacing_after=3
    )
    # Body: 11pt
    body_style = ensure_paragraph_style(
        doc,
        style_name='Resume Body',
        font_name='Calibri',
        font_size=11,
        spacing_before=0,
        spacing_after=6
    )
    # Bullets: 11pt
    bullet_style = ensure_paragraph_style(
        doc,
        style_name='List Bullet',
        font_name='Calibri',
        font_size=11,
        spacing_before=0,
        spacing_after=4
    )

    # 4) Main Heading (Name) in Document Body
    doc.add_paragraph(resume_data['heading']['name'], style=heading1_style)

    # 5) Biography
    bio_para = doc.add_paragraph(resume_data["biography"], style=body_style)
    bio_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # 6) Technology Skills
    doc.add_paragraph("Technology Skills", style=heading2_style)

    skills_table = doc.add_table(rows=0, cols=2)
    try:
        skills_table.style = 'Table Grid'
    except KeyError:
        skills_table.style = 'TableGrid'

    for key, value in resume_data["skill_set"].items():
        row_cells = skills_table.add_row().cells
        # Category cell
        category_run = row_cells[0].paragraphs[0].add_run(key)
        category_run.bold = True
        category_run.font.size = Pt(11)
        category_run.font.name = 'Calibri'

        # Detail cell
        detail_run = row_cells[1].paragraphs[0].add_run(value)
        detail_run.font.size = Pt(11)
        detail_run.font.name = 'Calibri'

    doc.add_paragraph("", style=body_style)  # spacing after table

    # 7) Work History
    doc.add_paragraph("Work History", style=heading2_style)

    for entry in resume_data["work_history"]:
        # Company & Location
        company_para = doc.add_paragraph(
            entry["company"], style=heading3_style)
        if entry.get("location"):
            company_para.add_run(f", {entry['location']}")

        # Position & Dates
        pos_date_par = doc.add_paragraph(
            f"{entry['position']} ({entry['dates']})",
            style=body_style
        )
        pos_date_par.paragraph_format.space_after = Pt(4)

        # Responsibilities with bullet style
        for responsibility in entry["responsibilities"]:
            resp_para = doc.add_paragraph(responsibility, style=bullet_style)
            resp_para.paragraph_format.left_indent = Inches(0.25)
            resp_para.paragraph_format.space_before = 0

    # 8) Save
    date_str = datetime.now().strftime('%Y%m%d')
    person_name = resume_data['heading']['name'].replace(' ', '_')
    doc_filename = f"{person_name}.{date_str}.docx"
    doc.save(doc_filename)

    return doc_filename


def ensure_paragraph_style(doc, style_name, font_name='Calibri', font_size=11,
                           bold=False, italic=False, spacing_before=0, spacing_after=0):
    """
    Ensures a paragraph style exists in the document with the given font properties.
    If it doesn't exist, it creates it; if it does, it updates it.
    Returns the style object for use.
    """
    styles = doc.styles
    if style_name in styles:
        style = styles[style_name]
    else:
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.font.bold = bold
    style.font.italic = italic

    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(spacing_before)
    paragraph_format.space_after = Pt(spacing_after)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    return style


# -------------------- EXAMPLE USAGE --------------------
if __name__ == "__main__":
      resume_data_example = {
       "heading": {
              "name": "Scott",
              "phone": "248-xxx-xxxx",
              "social_media": "https://github.com/scottsdevelopment/resume",
              "email": "scott@scottsdevelopment.com"
          },
          "biography": (
              "I’m a seasoned full-stack developer dedicated to making complex ideas "
              "simple and scalable. I focus on real-time automation, high-performance systems, and AI integration "
              "to deliver future-proof applications that bring real value to both "
              "developers and end-users."
          ),
          "skill_set": {
              "Operating Systems": "Linux, Windows, macOS",
              "Web Development": (
                  "HTML, CSS, JavaScript, TypeScript, Angular, React, Drupal, "
                  "WordPress, Shopify, WooCommerce, PHP"
              ),
              # ... Add additional skills as needed ...
          },
          "work_history": [
              {
                  "company": "Scott's Development",
                  "location": "Remote",
                  "position": "Chief Executive Officer",
                  "dates": "Oct 2024 – Present",
                  "responsibilities": [
                      "Spearheaded advanced automation for live-streaming services.",
                      "Designed and implemented scalable bot frameworks utilizing modern APIs."
                  ]
              }
              # ... Add more positions as needed ...
          ]
      }

try:
    synced_resume_path_2025 = generate_new_resume(resume_data_example)
    print(f"Synced and formatted resume saved to: {synced_resume_path_2025}")
except Exception as e:
    print(f"Error: {e}")
